#!/usr/bin/python
#-*-coding:utf-8-*-

import json

import pymysql
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor

from . import Column
from . import Table

__author__ = 'skydu'

class Mysql2docx(object):
    dbName = ''
    #
    def __init__(self):
        dbName = ''

    def getComment(self,comment):
        if comment==None:
            return ""
        try:
            data=json.loads(comment)
            return data[0]['value']
        except:
            return comment

    def getTables(self,db):
        sql = "select table_name, TABLE_COMMENT from information_schema.tables " \
              "where table_schema = '%s' and table_type = 'base table'"%self.dbName
        cursor=db.cursor()
        cursor.execute(sql)
        data = cursor.fetchall()
        tables=list()
        for table in data:
            t=Table(table[0], self.getComment(table[1]));
            tables.append(t)
        cursor.close()
        return tables

    def getColumns(self,db,tableName):
        sql = "SELECT  " \
              "COLUMN_NAME 列名,  " \
              "COLUMN_TYPE 数据类型,  " \
              "IF(IS_NULLABLE = 'NO', '√', '') 是否为空,    " \
              "COLUMN_DEFAULT 默认值,    " \
              "if(lower(COLUMN_NAME) = 'id', 'id', COLUMN_COMMENT) 备注  , " \
              "IF(COLUMN_KEY = 'PRI', '√', '') 主键 "\
              "FROM  INFORMATION_SCHEMA.COLUMNS  " \
              "where  table_schema ='%s'  AND   table_name  = '%s';" % (self.dbName, tableName)
        cursor = db.cursor()
        cursor.execute(sql)
        data = cursor.fetchall()
        columns=list()
        for column in data:
            c=Column(column[0],column[1],column[2],column[3],self.getComment(column[4]),column[5])
            columns.append(c)
        cursor.close()
        return columns

    @staticmethod
    def do(dbHost, dbUser, dbPassword, dbName, dbPort,doc='数据库设计文档.docx'):
        print("dbHost:%s,dbUser:%s,dbPassword:%s,dbName:%s,dbPort:%d" % (dbHost, dbUser, dbPassword, dbName, dbPort))
        instance=Mysql2docx()
        instance.dbName=dbName
        # db = pymysql.connect(dbHost, dbUser, dbPassword, dbName, dbPort, charset="utf8")
        db = pymysql.connect(host=dbHost, user=dbUser, password=dbPassword, database=dbName, port=dbPort, charset="utf8")

        tables = instance.getTables(db);
        for table in tables:
            tableName = table.name
            table.columns = instance.getColumns(db, tableName)

        document = Document()

        head2 = document.add_heading(level=2)
        run = head2.add_run(u'物理表')
        run.font.name=u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 段落后行距
        # head2.paragraph_format.space_after = Pt(30)


        for table in tables:
            print("table:%s" % table)
            # 三级标题
            head3 = document.add_heading(level=3)
            run = head3.add_run(u"表%s[%s]" % (table.name, table.comment))
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 段落后行距
            # head2.paragraph_format.space_after = Pt(30)

# Light Grid
# Light List Accent 1

            t = document.add_table(rows=len(table.columns) + 1, cols=8, style='Medium Grid 3 Accent 1')

            t.style.font.size = Pt(10.5)  # 字体大小15磅
            t.style.font.name = u'宋体'
            t.style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            t.style.font.color.rgb = RGBColor(0, 0, 0)
            t.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐

            cells = t.rows[0].cells
            cells[0].text = '#'
            cells[1].text = '字段'
            cells[2].text = '名称'
            cells[3].text = '数据类型'
            cells[4].text = '主键'
            cells[5].text = '非空'
            cells[6].text = '默认值'
            cells[7].text = '备注说明'
            i = 0
            for column in table.columns:
                i += 1
                rowCells = t.rows[i].cells
                rowCells[0].text = str(i)
                rowCells[1].text = column.name
                rowCells[2].text = column.comment
                rowCells[3].text = column.type
                rowCells[4].text = column.primaryKey
                rowCells[5].text = column.allowNull
                if column.defaultValue!=None:
                    rowCells[6].text = column.defaultValue
                # rowCells[7].text = column.comment
            #
            # document.add_page_break()
        #
        document.save(doc)
